import fitz  # PyMuPDF
import docx
import os
import io

def read_pdf(file_path: str) -> str:
    """Extracts text from a PDF file."""
    text = ""
    try:
        with fitz.open(file_path) as doc:
            for page in doc:
                text += page.get_text() + "\n"
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
    return text

def read_pdf_from_bytes(file_bytes: bytes) -> str:
    """Extracts text from a PDF file in memory."""
    text = "" 
    try:
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            for page in doc:
                text += page.get_text() + "\n"
    except Exception as e:
        print(f"Error reading PDF from bytes: {e}")
    return text

def read_docx(file_path: str) -> str:
    """Extracts text from a DOCX file."""
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
    return text

def read_docx_from_bytes(file_bytes: bytes) -> str:
    """Extracts text from a DOCX file in memory."""
    text = ""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        print(f"Error reading DOCX from bytes: {e}")
    return text

def read_resume(file_path: str) -> str:
    """Reads a resume file (PDF or DOCX) and returns its text content."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        return read_pdf(file_path)
    elif ext == '.docx':
        return read_docx(file_path)
    elif ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file format: {ext}. Only PDF, DOCX, and TXT are supported.")

def read_resume_from_bytes(file_bytes: bytes, filename: str) -> str:
    """Reads a resume file from memory and returns its text content."""
    ext = os.path.splitext(filename)[1].lower()
    if ext == '.pdf':
        return read_pdf_from_bytes(file_bytes)
    elif ext == '.docx':
        return read_docx_from_bytes(file_bytes)
    elif ext == '.txt':
        return file_bytes.decode('utf-8')
    else:
        raise ValueError(f"Unsupported file format: {ext}. Only PDF, DOCX, and TXT are supported.")
